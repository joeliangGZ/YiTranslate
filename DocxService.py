import asyncio
import os
import re
import time
import uuid
from io import BytesIO
from pathlib import Path
from typing import List, Dict

import uvicorn
from docx import Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse

from YiTranslateSolely import translate_text

app = FastAPI()

TEMPLATE_DIR = "templates/"
PRODUCT_DIR = "products/"
Path(TEMPLATE_DIR).mkdir(parents=True, exist_ok=True)
Path(PRODUCT_DIR).mkdir(parents=True, exist_ok=True)
CONCURRENCY_LIMIT = 20

class DocumentItem:
    def __init__(self, placeholder_number: int, original_content: str, translate_content: str = None):
        self.placeholder_number = placeholder_number
        self.original_content = original_content
        self.translate_content = translate_content

class DocumentEntity:
    def __init__(self, doc_id: str, original_filename: str, template_filename: str, items: List[DocumentItem]):
        self.doc_id = doc_id
        self.original_filename = original_filename
        self.template_filename = template_filename
        self.items = items

@app.post("/upload")
async def process_document(file: UploadFile = File(...)):
    # 1. 提取内容（假设已有实现）
    doc_entity = extract_content(file)

    # 2. 并发翻译
    await batch_translate(doc_entity.items)

    # 3. 填充模板（假设已有实现）
    return fill_template(doc_entity)

def extract_content(file: UploadFile):
    try:
        doc_id = str(uuid.uuid4())
        original_filename = file.filename or "uploaded.docx"
        base_name = original_filename.replace(".docx", "")
        template_filename = f"{base_name}_template_{int(time.time())}.docx"

        # 读取上传文件
        doc = Document(file.file)

        items = []
        current_num = 1

        # 处理段落
        for paragraph in doc.paragraphs:
            current_num = process_paragraph(paragraph, current_num, items)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_num = process_paragraph(paragraph, current_num, items)

        # 保存模板文件
        template_path = os.path.join(TEMPLATE_DIR, template_filename)
        doc.save(template_path)

        return DocumentEntity(doc_id, original_filename, template_filename, items)

    except Exception as e:
        raise HTTPException(status_code=500, detail="文档解析异常")

async def batch_translate(items: List[DocumentItem]):
    semaphore = asyncio.Semaphore(CONCURRENCY_LIMIT)
    tasks = [
        asyncio.create_task(limited_translate(item, semaphore))
        for item in items
    ]
    await asyncio.gather(*tasks)

async def limited_translate(item: DocumentItem, semaphore):
    async with semaphore:
        await translate_item(item)

async def translate_item(item: DocumentItem):
    try:
        # 核心调用点
        item.translate_content = await asyncio.to_thread(
            translate_text,
            item.original_content
        )
    except Exception as e:
        item.translate_content = f"[FAILED] {str(e)}"

def process_paragraph(paragraph: Paragraph, current_num: int, items: List[DocumentItem]) -> int:
    original_text = paragraph.text.strip()
    if not original_text:
        return current_num

    # 匹配编号模式
    pattern = re.compile(r"^(\d+(?:\.\d+)*\s+)(.*)$")
    match = pattern.match(original_text)

    if match:
        number_part = match.group(1)
        content = match.group(2)
        new_text = f"{number_part}{{{{ {current_num} }}}}"
    else:
        content = original_text
        new_text = f"{{{{ {current_num} }}}}"

    # 获取第一个 Run 的样式（如果存在）
    if paragraph.runs:
        source_run = paragraph.runs[0]
        font = source_run.font
        # 提取可复制的样式属性
        style = {
            "name": font.name,
            "size": font.size,
            "bold": font.bold,
            "italic": font.italic,
            "color": font.color.rgb
        }
    else:
        style = None

    # 清空段落内容
    paragraph.clear()

    # 创建新 Run 并应用样式
    new_run = paragraph.add_run(new_text)
    if style:
        new_run.font.name = style["name"]
        if style["size"]:
            new_run.font.size = Pt(style["size"])
        new_run.font.bold = style["bold"]
        new_run.font.italic = style["italic"]
        if style["color"]:
            new_run.font.color.rgb = style["color"]

    items.append(DocumentItem(current_num, content))
    return current_num + 1

def fill_template(doc_entity: DocumentEntity):
    try:
        template_path = os.path.join(TEMPLATE_DIR, doc_entity.template_filename)
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="模板文件未找到")

        # Using attribute access for items as well
        content_map = {item.placeholder_number: item.translate_content
                       for item in doc_entity.items}

        doc = Document(template_path)

        # Process paragraphs
        for paragraph in doc.paragraphs:
            process_fill_paragraph(paragraph, content_map)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_fill_paragraph(paragraph, content_map)

        # 将文档保存到内存流
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # 将内存中的内容写入到本地文件
        output_filename = doc_entity.original_filename
        output_path = os.path.join(PRODUCT_DIR, output_filename)
        with open(output_path, "wb") as f:
            f.write(buffer.getvalue())

        # 为了防止后面流指针位置问题，再创建一个内存流返回下载
        download_stream = BytesIO(buffer.getvalue())
        download_stream.seek(0)

        return StreamingResponse(
            download_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={output_filename}"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail="文档填充异常")

def process_fill_paragraph(paragraph: Paragraph, content_map: Dict[int, str]):
    original_text = paragraph.text.strip()
    if not original_text:
        return

    # 匹配占位符
    pattern = re.compile(r"\{\{\s*(\d+)\s*\}\}")
    new_text = pattern.sub(lambda m: content_map.get(int(m.group(1)), m.group(0)), original_text)

    # 获取第一个 Run 的样式（如果存在）
    if paragraph.runs:
        source_run = paragraph.runs[0]
        font = source_run.font
        style = {
            "name": font.name,
            "size": font.size,
            "bold": font.bold,
            "italic": font.italic,
            "color": font.color.rgb
        }
    else:
        style = None

    # 清空段落内容
    paragraph.clear()

    # 创建新 Run 并应用样式
    new_run = paragraph.add_run(new_text)
    if style:
        new_run.font.name = style["name"]
        if style["size"]:
            new_run.font.size = Pt(style["size"])
        new_run.font.bold = style["bold"]
        new_run.font.italic = style["italic"]
        if style["color"]:
            new_run.font.color.rgb = style["color"]

if __name__ == "__main__":
    uvicorn.run("DocxService:app", host="0.0.0.0", port=8080)
