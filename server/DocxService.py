import asyncio
import os
import re
import time
import uuid
from io import BytesIO
from pathlib import Path
from typing import List

import uvicorn
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse

from YiTranslateSolely import translate_text

app = FastAPI()

TEMPLATE_DIR = "templates/"
PRODUCT_DIR = "products/"
Path(TEMPLATE_DIR).mkdir(parents=True, exist_ok=True)
Path(PRODUCT_DIR).mkdir(parents=True, exist_ok=True)
CONCURRENCY_LIMIT = 100

class DocumentItem:
    def __init__(self, placeholder_number: int, original_content: str, translate_content: str = None):
        self.placeholder_number = placeholder_number
        self.original_content = original_content
        self.translate_content = translate_content

class DocumentEntity:
    def __init__(self, doc_id: str, original_filename: str, template_filename: str, items: List[DocumentItem]):
        self.doc_id = doc_id
        self.original_filename = original_filename
        self.template_filename = template_filename
        self.items = items

@app.post("/upload")
async def process_document(file: UploadFile = File(...)):

    doc_entity = extract_content(file)

    await batch_translate(doc_entity.items)
    return fill_template(doc_entity)


def extract_content(file: UploadFile) -> DocumentEntity:
    try:
        doc_id = str(uuid.uuid4())
        original_filename = file.filename or "uploaded.docx"
        base_name = os.path.splitext(original_filename)[0]
        template_filename = f"{base_name}_template_{int(time.time())}.docx"

        # 读取上传文件到内存
        file_content = file.file.read()
        file_stream = BytesIO(file_content)
        doc = Document(file_stream)

        items: List[DocumentItem] = []
        current_num = 1

        # 处理段落
        for paragraph in doc.paragraphs:
            current_num = process_paragraph(paragraph, current_num, items)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_num = process_paragraph(paragraph, current_num, items)

        # 保存模板文件 —— 使用 python-docx 的 save 接口
        template_path = os.path.join(TEMPLATE_DIR, template_filename)
        doc.save(template_path)

        # 重置文件流指针（以防后续重用）
        file.file.seek(0)

        return DocumentEntity(doc_id, original_filename, template_filename, items)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档解析异常：{e}")


def clear_paragraph(paragraph: Paragraph):
    """
    安全地清空段落中的所有 run，避免破坏底层 XML 结构。
    """
    for run in paragraph.runs[::-1]:
        paragraph._element.remove(run._element)

def process_paragraph(paragraph: Paragraph, current_num: int, items: List[DocumentItem]) -> int:
    original_text = paragraph.text.strip()
    if not original_text:
        return current_num

    # 匹配编号模式，比如 "1.2. 文本内容"
    pattern = re.compile(r"^(\d+(?:\.\d+)*\s+)(.*)$")
    match = pattern.match(original_text)

    if match:
        number_part = match.group(1)
        content = match.group(2)
        new_text = f"{number_part}{{{{ {current_num} }}}}"
    else:
        content = original_text
        new_text = f"{{{{ {current_num} }}}}"

    # 尝试保留第一个 run 的样式
    style = None
    if paragraph.runs:
        run0 = paragraph.runs[0]
        font = run0.font
        style = {
            "name": font.name or "SimSun",  # 默认中文宋体
            "size": font.size or Pt(12),
            "bold": font.bold,
            "italic": font.italic,
            "color": font.color.rgb
        }

    # 清空段落
    clear_paragraph(paragraph)

    # 添加新的 run
    new_run = paragraph.add_run(new_text)
    if style:
        # 设置西文字体
        new_run.font.name = style["name"]
        # 设置东亚文字（中文）字体
        new_run._element.rPr.rFonts.set(qn('w:eastAsia'), style["name"])
        # 字号
        new_run.font.size = style["size"]
        # 加粗、斜体
        new_run.font.bold = style["bold"]
        new_run.font.italic = style["italic"]
        # 颜色
        if style["color"]:
            new_run.font.color.rgb = style["color"]

    items.append(DocumentItem(current_num, content))
    return current_num + 1


async def batch_translate(items: List[DocumentItem]):
    """
    并发调用你的 translate_text 函数，把 items 中的 original_content 翻译后填充到 translate_content。
    """
    semaphore = asyncio.Semaphore(CONCURRENCY_LIMIT)

    async def _translate(item: DocumentItem):
        async with semaphore:
            item.translate_content = await translate_text(item.original_content)

    await asyncio.gather(*[_translate(item) for item in items])

def fill_template(doc_entity: DocumentEntity):
    """
    读取模板文件，替换 {{ n }} 占位符为翻译内容，并返回下载流。
    """
    template_path = os.path.join(TEMPLATE_DIR, doc_entity.template_filename)
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for item in doc_entity.items:
            placeholder = f"{{{{ {item.placeholder_number} }}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, item.translate_content or "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for item in doc_entity.items:
                    placeholder = f"{{{{ {item.placeholder_number} }}}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, item.translate_content or "")

    # 输出最终文档
    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    filename = f"translated_{doc_entity.original_filename}"
    return StreamingResponse(
        output_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

if __name__ == "__main__":
    uvicorn.run("DocxService:app", host="0.0.0.0", port=8081)
